import os
import re
import math
import hashlib
import json
import base64
import logging
from typing import List, Dict, Optional, Tuple
from fastapi import HTTPException, status, UploadFile
from app.config import settings
import requests

# Nuevos importes para Memoria Híbrida
from app.services.openrouter_embeddings import openrouter_embeddings
from app.services.vector_index import vector_index

logger = logging.getLogger("document_engine")

# Límites de seguridad y costos para OCR
MAX_FILE_SIZE_BYTES = 25 * 1024 * 1024  # 25 MB
MAX_OCR_PAGES_COUNT = 25  # Límite estricto SOLO cuando se requiere OCR con IA para evitar consumo masivo de tokens

class DocumentChunk:
    def __init__(self, chunk_index: int, text: str, page_number: int = 1, metadata: dict = None):
        self.chunk_index = chunk_index
        self.text = text
        self.page_number = page_number
        self.metadata = metadata or {}

class ProcessedDocument:
    def __init__(self, doc_id: str, user_id: int, filename: str, doc_type: str, file_hash: str, pages_count: int):
        self.doc_id = doc_id
        self.user_id = user_id
        self.filename = filename
        self.doc_type = doc_type
        self.file_hash = file_hash
        self.pages_count = pages_count
        self.chunks: List[DocumentChunk] = []
        self.created_at = "2026" # Timestamp o referencia

    def to_dict(self):
        return {
            "doc_id": self.doc_id,
            "user_id": self.user_id,
            "filename": self.filename,
            "doc_type": self.doc_type,
            "file_hash": self.file_hash,
            "pages_count": self.pages_count,
            "chunks": [
                {"chunk_index": c.chunk_index, "text": c.text, "page_number": c.page_number, "metadata": c.metadata}
                for c in self.chunks
            ]
        }

class DocumentRAGEngine:
    def __init__(self):
        # Memoria / caché persistente de documentos por doc_id
        self.cache_dir = os.path.join(os.path.abspath(os.path.dirname(__file__)), "../../rag_documents_cache")
        if not os.path.exists(self.cache_dir):
            os.makedirs(self.cache_dir, exist_ok=True)
        self.documents: Dict[str, ProcessedDocument] = {}
        self._load_cache_from_disk()

    def _load_cache_from_disk(self):
        try:
            cache_file = os.path.join(self.cache_dir, "documents_registry.json")
            if os.path.exists(cache_file):
                with open(cache_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    for doc_id, item in data.items():
                        doc = ProcessedDocument(
                            doc_id=item["doc_id"],
                            user_id=item["user_id"],
                            filename=item["filename"],
                            doc_type=item.get("doc_type", "document"),
                            file_hash=item["file_hash"],
                            pages_count=item.get("pages_count", 1)
                        )
                        for c in item.get("chunks", []):
                            doc.chunks.append(DocumentChunk(
                                chunk_index=c["chunk_index"],
                                text=c["text"],
                                page_number=c.get("page_number", 1),
                                metadata=c.get("metadata", {})
                            ))
                        self.documents[doc_id] = doc
                logger.info(f"RAG Document Engine cache cargado: {len(self.documents)} documentos registrados.")
        except Exception as e:
            logger.error(f"Error cargando registro RAG cache: {e}")

    def _save_cache_to_disk(self):
        try:
            cache_file = os.path.join(self.cache_dir, "documents_registry.json")
            data = {doc_id: doc.to_dict() for doc_id, doc in self.documents.items()}
            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"Error guardando registro RAG cache: {e}")

    def _perform_fast_ocr(self, file_bytes: bytes, mime_type: str = "application/pdf") -> str:
        """
        Utiliza Gemini 2.0 Flash Lite (el modelo multimodal más ágil, económico y veloz en OpenRouter)
        para realizar OCR clínico impecable sobre archivos escaneados.
        """
        if not settings.OPENROUTER_API_KEY:
            return "[Texto escaneado - OCR no disponible por falta de API Key]"
            
        try:
            logger.info("Activando OCR de alta velocidad con Gemini 2.0 Flash Lite en OpenRouter...")
            b64_data = base64.b64encode(file_bytes).decode("utf-8")
            
            headers = {
                "Authorization": f"Bearer {settings.OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
                "HTTP-Referer": "http://localhost:3000",
                "X-Title": "SUPER-UCE DOC - OCR Engine"
            }
            
            prompt_instructions = (
                "Actúa como un motor de reconocimiento óptico (OCR) clínico especializado y de precisión absoluta. "
                "Extrae y transcribe literalmente todo el texto médico, tablas, nombres de medicamentos, dosis y resultados "
                "de laboratorio legibles en el documento/imagen adjunta en español. NUNCA resumas ni des explicaciones, "
                "únicamente devuelve el contenido literal extraído de forma limpia y ordenada."
            )
            
            # Modelo más ágil y económico para OCR y transcripciones visuales
            ocr_model = "google/gemini-2.5-flash"
            
            payload = {
                "model": ocr_model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt_instructions},
                            {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{b64_data}"}}
                        ]
                    }
                ],
                "temperature": 0.1,
                "max_tokens": 2500
            }
            
            res = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload, timeout=20)
            if res.status_code == 200:
                extracted = res.json()["choices"][0]["message"]["content"].strip()
                return re.sub(r'^(Transcripción(:|\s)|Respuesta:)\s*', '', extracted, flags=re.IGNORECASE).strip(' "')
            else:
                logger.warning(f"Fallo en OCR con Gemini 2.0 Flash Lite (HTTP {res.status_code}): {res.text}")
        except Exception as e:
            logger.error(f"Error durante OCR ágil: {e}")
            
        return "[Error al procesar OCR sobre archivo escaneado]"

    def _clean_and_normalize_text(self, text: str) -> str:
        """
        Elimina espacios innecesarios, encabezados/pies repetitivos, páginas en blanco y normaliza.
        """
        if not text:
            return ""
            
        # 1. Eliminar numeraciones de página repetitivas ("Página X de Y", "Page X", etc.)
        cleaned = re.sub(r'(?i)\b(página|pag|page)\s*\d+\s*(de|of)?\s*\d*\b', '', text)
        cleaned = re.sub(r'(\n\s*[-–—]\s*\d+\s*[-–—]\s*\n)', '\n', cleaned)
        
        # 2. Reemplazar múltiples saltos de línea (>2) por un doble salto estándar (párrafos)
        cleaned = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned)
        
        # 3. Eliminar espacios en blanco excesivos por línea
        lines = [line.strip() for line in cleaned.splitlines()]
        # Filtrar líneas vacías repetitivas y conservar contenido
        filtered_lines = []
        last_was_empty = False
        for line in lines:
            if not line:
                if not last_was_empty:
                    filtered_lines.append("")
                last_was_empty = True
            else:
                filtered_lines.append(line)
                last_was_empty = False
                
        result = "\n".join(filtered_lines).strip()
        return result

    def _smart_chunking(self, text: str, max_tokens_approx: int = 950) -> List[str]:
        """
        División automática en chunks de aprox. 800 a 1200 tokens (~3000 a 4500 caracteres),
        respetando fronteras de párrafos y secciones clínicas para no cortar ideas en la mitad.
        """
        max_chars = max_tokens_approx * 4  # ~3800 caracteres
        min_chars = 400
        
        # Primero separamos por párrafos dobles o encabezados en mayúsculas
        paragraphs = re.split(r'\n\s*\n', text)
        chunks = []
        current_chunk = ""
        
        for p in paragraphs:
            p_clean = p.strip()
            if not p_clean:
                continue
                
            if len(current_chunk) + len(p_clean) + 2 <= max_chars:
                current_chunk += ("\n\n" + p_clean if current_chunk else p_clean)
            else:
                # Si el bloque actual supera el mínimo operacional de tamaño, lo guardamos
                if len(current_chunk) >= min_chars or not chunks:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = p_clean
                else:
                    current_chunk += "\n\n" + p_clean
                    
        if current_chunk:
            chunks.append(current_chunk)
            
        # Si un solo párrafo excedió enormemente max_chars, lo separamos por puntos seguidos (. )
        final_chunks = []
        for c in chunks:
            if len(c) > max_chars * 1.3:
                sentences = re.split(r'(?<=\.)\s+', c)
                sub_chunk = ""
                for s in sentences:
                    if len(sub_chunk) + len(s) + 1 <= max_chars:
                        sub_chunk += (" " + s if sub_chunk else s)
                    else:
                        if sub_chunk:
                            final_chunks.append(sub_chunk)
                        sub_chunk = s
                if sub_chunk:
                    final_chunks.append(sub_chunk)
            else:
                final_chunks.append(c)
                
        # Deduplicación exacta en bloques
        unique_chunks = []
        seen = set()
        for fc in final_chunks:
            normalized_compare = re.sub(r'\s+', '', fc.lower())
            if normalized_compare not in seen and len(fc) > 15:
                seen.add(normalized_compare)
                unique_chunks.append(fc)
                
        return unique_chunks

    def parse_and_index_document(self, file_bytes: bytes, filename: str, user_id: int = 1) -> ProcessedDocument:
        """
        Procesa un archivo entrante, aplica límites, extrae texto, limpia, hace chunking y lo almacena en caché.
        """
        # 1. Verificar tamaño máximo de archivo
        if len(file_bytes) > MAX_FILE_SIZE_BYTES:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"El archivo excede el tamaño máximo permitido de 25 MB ({len(file_bytes) // (1024*1024)} MB recibidos)."
            )

        # 2. Calcular SHA-256 para verificación de Caché (Deduplicación)
        file_hash = hashlib.sha256(file_bytes).hexdigest()
        
        # Si ya existe para este usuario y no ha cambiado, devolvemos al instante
        for doc in self.documents.values():
            if doc.user_id == user_id and doc.file_hash == file_hash:
                logger.info(f"Caché RAG golpeado para archivo: {filename} (Hash: {file_hash[:8]}...). Reutilizando {len(doc.chunks)} chunks.")
                return doc

        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "txt"
        extracted_text = ""
        pages_count = 1
        doc_type = ext.upper()
        
        # 3. Extracción por tipo de archivo
        if ext == "txt":
            try:
                extracted_text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                extracted_text = file_bytes.decode("latin-1", errors="ignore")
                
        elif ext == "pdf":
            import io
            try:
                import pypdf
                pdf_reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                pages_count = len(pdf_reader.pages)
                page_texts = []
                for idx, page in enumerate(pdf_reader.pages):
                    pt = page.extract_text() or ""
                    if len(pt.strip()) > 0:
                        page_texts.append(f"[Página {idx+1}]\n{pt}")
                        
                extracted_text = "\n\n".join(page_texts).strip()
                # Si el PDF entero tiene menos de 50 caracteres, significa que es un escaneo de fotos/imágenes (requiere OCR con IA)
                if len(extracted_text) < 50:
                    if pages_count > MAX_OCR_PAGES_COUNT:
                        raise HTTPException(
                            status_code=status.HTTP_400_BAD_REQUEST,
                            detail=f"El archivo escaneado (sin capa de texto) excede el límite de {MAX_OCR_PAGES_COUNT} páginas para OCR mediante Inteligencia Artificial ({pages_count} págs detectadas). Esto se implementa para evitar desbordes en el consumo de tokens. Los archivos PDF con texto seleccionable o ficheros normales no tienen límite de páginas ni de palabras."
                        )
                    logger.info("PDF escaneado o de imagen detectado (< 50 caracteres). Activando OCR...")
                    extracted_text = self._perform_fast_ocr(file_bytes, "application/pdf")
            except HTTPException as he:
                raise he
            except Exception as e:
                if pages_count > MAX_OCR_PAGES_COUNT:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"El archivo no pudo leerse nativamente y supera las {MAX_OCR_PAGES_COUNT} páginas permitidas para escaneo inteligente con IA ({pages_count} págs)."
                    )
                logger.error(f"Error extrayendo PDF con pypdf, aplicando OCR fallback: {e}")
                extracted_text = self._perform_fast_ocr(file_bytes, "application/pdf")
                
        elif ext in ["docx", "doc"]:
            import io
            try:
                import docx
                doc_stream = io.BytesIO(file_bytes)
                docx_doc = docx.Document(doc_stream)
                paras = [p.text for p in docx_doc.paragraphs if p.text.strip()]
                # También extraer contenido de tablas médicas o de laboratorio
                for table in docx_doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            paras.append(row_text)
                extracted_text = "\n\n".join(paras)
            except Exception as e:
                logger.error(f"Error procesando DOCX: {e}")
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="No se pudo interpretar el archivo DOCX. Verifique que no esté dañado o protegido."
                )
                
        elif ext in ["png", "jpg", "jpeg", "webp"]:
            logger.info(f"Imagen detectada (.{ext}). Activando OCR...")
            mime_type = f"image/{'jpeg' if ext == 'jpg' else ext}"
            extracted_text = self._perform_fast_ocr(file_bytes, mime_type)
            
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Formato no soportado (.{ext}). Suba archivos en PDF, DOCX, TXT, o imágenes (PNG, JPG, WEBP)."
            )

        # 4. No aplicamos límite global de palabras/páginas para texto nativo, RAG se encarga en memoria.
        logger.info(f"Texto extraído del archivo {filename}: ~{len(extracted_text.split())} palabras en {pages_count} págs.")
        # 5. Limpieza y Normalización
        clean_text = self._clean_and_normalize_text(extracted_text)
        if not clean_text or len(clean_text) < 10:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="El documento no contiene texto legible ni datos clínicos útiles tras su procesamiento."
            )

        # 6. Chunking semántico
        chunks_text = self._smart_chunking(clean_text)
        doc_id = f"doc_{file_hash[:12]}"
        
        proc_doc = ProcessedDocument(
            doc_id=doc_id,
            user_id=user_id,
            filename=filename,
            doc_type=doc_type,
            file_hash=file_hash,
            pages_count=pages_count
        )
        
        # Batch fetching for embeddings drastically reduces network time
        embeddings = openrouter_embeddings.get_embeddings(chunks_text)
        
        for idx, ct in enumerate(chunks_text):
            # Detectar número de página probable en el texto
            page_num = 1
            pm = re.search(r'\[Página\s+(\d+)\]', ct, flags=re.IGNORECASE)
            if pm:
                page_num = int(pm.group(1))
            proc_doc.chunks.append(DocumentChunk(chunk_index=idx+1, text=ct, page_number=page_num))
            
            # Indexación en Memoria Híbrida Vectorial
            if idx < len(embeddings):
                emb = embeddings[idx]
                if emb:
                    vector_index.add_record(
                        embedding=emb,
                        payload={
                            "type": "rag_chunk",
                            "doc_id": doc_id,
                            "text": ct,
                            "page_number": page_num,
                            "filename": filename
                        }
                    )
                
        self.documents[doc_id] = proc_doc
        self._save_cache_to_disk()
        logger.info(f"Documento procesado exitosamente: {filename} -> ID {doc_id} con {len(proc_doc.chunks)} chunks RAG.")
        return proc_doc

    def get_document(self, doc_id: str) -> Optional[ProcessedDocument]:
        return self.documents.get(doc_id)

    def search_hybrid_chunks(self, doc_ids: List[str], query: str, top_k: int = 6) -> List[Tuple[str, str, float]]:
        """
        Búsqueda 100% Semántica impulsada por text-embedding-004 y Vector Index.
        Reemplaza la lógica anterior basada en cruce de palabras clave (TF-IDF).
        """
        target_docs = [self.documents[d_id] for d_id in doc_ids if d_id in self.documents]
        if not target_docs:
            return []
            
        query_clean = query.strip()
        if not query_clean or len(query_clean) < 5:
            # Fallback trivial si la query es vacía, devolver primeros chunks
            fallback = []
            for doc in target_docs:
                for ch in doc.chunks[:top_k]:
                    fallback.append((ch.text, f"[{doc.filename} - Pág {ch.page_number}]", 1.0))
            return fallback[:top_k]
            
        # Generar vector para la query
        q_emb = openrouter_embeddings.get_embedding(query_clean)
        if not q_emb:
            return []
            
        # Buscar en el vector index, filtrando sólo los doc_ids permitidos
        # Extraemos todos los chunks del vector que correspondan a estos documentos
        best_chunks = []
        for doc_id in doc_ids:
            results = vector_index.search(q_emb, filter_dict={"type": "rag_chunk", "doc_id": doc_id}, top_k=top_k)
            for score, meta in results:
                src_label = f"[{meta.get('filename', 'Doc')} - Pág {meta.get('page_number', 1)}]"
                best_chunks.append((score, meta.get("text", ""), src_label))
                
        best_chunks.sort(key=lambda x: x[0], reverse=True)
        
        # Formato de retorno: List[(text, source, score)]
        return [(txt, src, sc) for sc, txt, src in best_chunks[:top_k]]

# Singleton para uso en toda la aplicación
document_engine = DocumentRAGEngine()
